from flask import make_response, jsonify

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import numpy as np
import io
import datetime

from app.models.ExcelData import ExcelData

# Generates a Document object containing all the info for the Syllabus
def generateSyllabus(professor, course, CRN):
    # Compile time settings
    # Table style setting; this table style
    table_style = 'Light Shading Accent 1' #blue alternating
    #table_style = 'Light Shading' # grey alternating

    doc = Document()

    # add metadata
    doc.core_properties.author = professor.user_flashlineID
    doc.core_properties.language = 'English'
    doc.core_properties.title = course.course_name + ' ' + \
        course.course_semester + ' ' + course_year + ' Syllabus'
    doc.core_properties.comments = 'Made with the waffle iron'
    doc.core_properties.category = 'Syllabus'
    doc.core_properties.created = datetime.datetime.now()
    doc.core_properties.modifier = datetime.datetime.now()
    doc.core_properties.identifier = CRN.CRN

    doc.add_heading(course.course_name, 0)
    doc.add_heading('CRN: ' + CRN.CRN + ' - ' + \
            course.course_semester + ' ' + course_year, 1)

    doc.add_paragraph('\n')

    contact_table = doc.add_table(rows=len(user_contactfields), cols=2, \
            style=table_style)
    contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # contact table
    for i in range(0, len(professor.user_contactfields)):
        row = contact_table.rows[i].cells
        p = row[0].add_paragraph(user_contactfields[i][0])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = row[1].add_paragraph(user_contactfields[i][1])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # course fields
    for i in range(0, len(course.course_fields)):
        doc.add_heading(course.course_fields[i][0], level=1)
        doc.add_paragraph(course_fields[i][1])

    return doc

def docToBase64(doc):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return b64decode(buffer.getvalue())


# Parse data out of an excel file
# Takes an io.BytesIO type
def parseExcelFile(excel_file):
    rows_to_skip = 0
    wb = pd.read_excel(excel_file)

    # the beginning of the header of the file has more than two rows, therefore
    # we can reason the second row will not be NaN. Therefore, we can reason
    # that the first row with data there is the header row
    while not isinstance(wb.iloc[rows_to_skip][1], str):
        wb.drop(index=rows_to_skip)
        rows_to_skip += 1

    # we add 1 to rows_to_skip as the top row is automatically set as the name
    # of the column, and therefore not indexable the first time we read the
    # file
    rows_to_skip += 1

    #re import the document, skipping the correct rows
    wb = pd.read_excel(excel_file, skiprows=rows_to_skip)

    #drop last row, if its the date the file was generated on
    if isinstance(wb.tail(1).iloc[0][0], datetime.datetime):
        wb.drop(wb.tail(1).index, inplace=True)

    # Dictionary to return
    ret = {}
    prev_crn = 0

    #TODO extract data
    for row in wb.index:
        # CASE 1: This is a continued line, we should populate the previous
        # CRN's values
        if np.isnan(wb["CRN"][row]):
            # set row, eliminate NaN values
            currRow = wb.iloc[row]
            currRow = currRow.fillna('')

            # update meeting places
            if not currRow["Bldg"] == '':
                ret[prev_crn].multipleMeetingPlaces = True
                ret[prev_crn].building += ";" + wb["Bldg"][row].strip()
            if not currRow["Room"] == '':
                ret[prev_crn].multipleMeetingPlaces = True
                ret[prev_crn].room += ";" + wb["Room"][row]

            # update meeting time
            if not currRow["Times"] == '':
                ret[prev_crn].multipleMeetingTimes = True
                ret[prev_crn].time += ";" + wb["Times"][row].strip()

            # update meeting days
            if not currRow["Meeting Days"] == '':
                ret[prev_crn].multipleMeetingDays = True
                ret[prev_crn].meetingDays += ";" + wb["Meeting Days"][row].strip()

        # CASE 2: We have a CRN, then populate new entry
        else:
            # set prev CRN
            prev_crn = wb["CRN"][row]

            # set row, eliminate NaN values
            currRow = wb.iloc[row]
            currRow = currRow.fillna('')

            # populate ret dict with ExcelData object, then populate data into
            # it
            ret[prev_crn] = ExcelData()

            ret[prev_crn].CRN = prev_crn

            ret[prev_crn].courseNumber    = ('%f' % currRow["Course#"]).rstrip('0').rstrip('.')
            ret[prev_crn].section         = currRow["Section"]
            ret[prev_crn].title           = currRow["Title"].strip()
            ret[prev_crn].instructorEmail = currRow["Instructor Email Address"].strip()
            ret[prev_crn].building        = currRow["Bldg"].strip()
            ret[prev_crn].room            = currRow["Room"]
            ret[prev_crn].time            = currRow["Times"].strip()
            ret[prev_crn].meetingDays     = currRow["Meeting Days"].strip()
            ret[prev_crn].instructorName  = currRow["Instructor"].strip()

    return ret

# Jsonifies the passed object, and makes a response object out of it
def sendResponse(result):
    resp = make_response(jsonify(result))
    resp.mimetype = 'application/json'
    return resp

# Returns the current year
def getYear():
    return datetime.datetime.today().year

# Returns the current semester as a string
def getSemester():
    today = datetime.datetime.today()
    if 1 <= today.month <= 5:
        return 'SPRING'
    elif 6 <= today.month < 8:
        return 'SUMMER'
    else:
        return 'FALL'
